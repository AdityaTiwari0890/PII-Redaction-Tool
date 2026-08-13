import copy
import re
from typing import Generator

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

from redactor.pipeline import PIIPipeline, RedactionConfig
from redactor.pseudonyms import DeterministicFaker


def _para_flat_text(para: Paragraph):
    flat = ""
    run_map = []
    for run in para.runs:
        start = len(flat)
        flat += run.text
        run_map.append((start, len(flat), run))
    return flat, run_map


def _apply_replacements_to_para(
    para: Paragraph,
    flat_text: str,
    run_map: list,
    replacements: list,
) -> None:
    if not replacements:
        return

    new_text = flat_text
    for start, end, repl in sorted(replacements, key=lambda x: x[0], reverse=True):
        new_text = new_text[:start] + repl + new_text[end:]

    if not para.runs:
        para.add_run(new_text)
        return

    first_run = para.runs[0]
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)
    first_run.text = new_text


def _process_paragraph(
    para: Paragraph,
    pipeline: PIIPipeline,
    dfaker: DeterministicFaker,
    mapping_log: list,
) -> None:
    flat, run_map = _para_flat_text(para)
    if not flat.strip():
        return

    redacted, log = pipeline.redact_text(flat, dfaker)
    if log:
        spans = pipeline.detect_all(flat)
        reps_from_spans = []
        for s in spans:
            fake_val = dfaker.get(s.text, s.label)
            reps_from_spans.append((s.start, s.end, fake_val))

        _apply_replacements_to_para(para, flat, run_map, reps_from_spans)
        mapping_log.extend(log)


def _process_cell(
    cell: _Cell,
    pipeline: PIIPipeline,
    dfaker: DeterministicFaker,
    mapping_log: list,
) -> None:
    for para in cell.paragraphs:
        _process_paragraph(para, pipeline, dfaker, mapping_log)


def process_docx(
    input_path: str,
    output_path: str,
    pipeline: PIIPipeline,
    dfaker: DeterministicFaker,
) -> list:
    doc = Document(input_path)
    mapping_log = []

    total_para = len(doc.paragraphs)
    for idx, para in enumerate(doc.paragraphs, 1):
        if idx % 100 == 0 or idx == total_para:
            print(f"  [paragraphs] {idx}/{total_para}", end="\r", flush=True)
        _process_paragraph(para, pipeline, dfaker, mapping_log)

    print()

    for ti, table in enumerate(doc.tables, 1):
        for row in table.rows:
            for cell in row.cells:
                _process_cell(cell, pipeline, dfaker, mapping_log)

    doc.save(output_path)
    return mapping_log


def iter_all_text_blocks(doc: Document) -> Generator[str, None, None]:
    for para in doc.paragraphs:
        flat = "".join(r.text for r in para.runs)
        if flat.strip():
            yield flat

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    yield txt
