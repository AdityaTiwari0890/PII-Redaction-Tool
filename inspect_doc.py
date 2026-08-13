"""Temporary inspection script - run once then delete."""
from docx import Document
import re

doc = Document(r'docs\Red Herring Prospectus.docx')
texts = [p.text for p in doc.paragraphs if p.text.strip()]
print(f"Total paragraphs: {len(texts)}")
print("=== FIRST 100 paragraphs ===")
for i, t in enumerate(texts[:100]):
    print(f"[{i}] {repr(t[:300])}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows[:4]):
        for ci, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"  [{ri},{ci}] {repr(cell.text[:250])}")

# Look for PII patterns
print("\n=== PII SCAN ===")
full_text = '\n'.join(p.text for p in doc.paragraphs)
emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', full_text)
phones = re.findall(r'(?:\+91[\s\-]?)?\d{5}[\s\-]?\d{5}', full_text)
pans   = re.findall(r'\b[A-Z]{5}\d{4}[A-Z]\b', full_text)
cins   = re.findall(r'\b[A-Z]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6}\b', full_text)
pins   = re.findall(r'\b\d{6}\b', full_text)
print(f"Emails found: {emails[:10]}")
print(f"Phones found: {phones[:10]}")
print(f"PANs found:   {pans[:10]}")
print(f"CINs found:   {cins[:5]}")
print(f"6-digit pins: {pins[:10]}")
print(f"\nTotal paragraphs: {len(texts)}, Total tables: {len(doc.tables)}")
