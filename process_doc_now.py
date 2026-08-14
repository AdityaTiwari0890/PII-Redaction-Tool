import os
import sys
import re
import hashlib
from pathlib import Path

try:
    from docx import Document
    from faker import Faker
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "faker", "-q"])
    from docx import Document
    from faker import Faker

ROOT = Path(__file__).resolve().parent
INPUT_PATH = ROOT / "docs" / "Red Herring Prospectus.docx"
OUTPUT_PATH = ROOT / "output" / "redacted.docx"

OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

fake = Faker("en_IN")
fake.seed_instance(42)

POOLS = {
    "PERSON": [fake.name() for _ in range(500)],
    "EMAIL": [fake.email() for _ in range(500)],
    "PHONE": [f"+91 {fake.msisdn()[3:]}" for _ in range(300)],
    "ADDRESS": [fake.address().replace("\n", ", ") for _ in range(200)],
    "PAN": [fake.bothify("?????#####?", letters="ABCDEFGHIJKLMNOPQRSTUVWXYZ") for _ in range(100)],
    "CIN": [fake.bothify("U#####??####PLC######", letters="ABCDEFGHIJKLMNOPQRSTUVWXYZ") for _ in range(100)],
    "IP_ADDRESS": [fake.ipv4() for _ in range(100)],
    "SSN": [fake.ssn() for _ in range(100)],
    "CREDIT_CARD": [fake.credit_card_number("visa") for _ in range(100)],
    "DOB": [fake.date_of_birth(minimum_age=25, maximum_age=65).strftime("%B %d, %Y") for _ in range(100)],
}

CACHE = {}

def get_fake(original, label):
    key = (original.strip(), label)
    if key not in CACHE:
        pool = POOLS.get(label, ["[REDACTED]"])
        idx = int(hashlib.md5((original + label).encode()).hexdigest(), 16) % len(pool)
        CACHE[key] = pool[idx]
    return CACHE[key]

_PERSONAL_DOMAINS = {"gmail.com", "yahoo.com", "yahoo.co.in", "hotmail.com", "outlook.com", "live.com", "rediffmail.com", "icloud.com"}
_SKIP_EXT = re.compile(r"\.(jpeg|png|jpg|gif|pdf|docx|xlsx)$", re.I)

PATTERNS = [
    (re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"), "EMAIL"),
    (re.compile(r"\+91[\s\-]?\d{5}[\s\-]?\d{5}"), "PHONE"),
    (re.compile(r"\b[6-9]\d{4}[\s\-]?\d{5}\b"), "PHONE"),
    (re.compile(r"\b[A-Z]{5}\d{4}[A-Z]\b"), "PAN"),
    (re.compile(r"\b[A-Z]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6}\b"), "CIN"),
    (re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b"), "IP_ADDRESS"),
    (re.compile(r"\b(Mr|Mrs|Ms|Dr|Prof|Shri|Smt|Sri)\.?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b"), "PERSON"),
]

def redact_string(text):
    if not text.strip():
        return text
    modified = text
    for pat, label in PATTERNS:
        for m in reversed(list(pat.finditer(modified))):
            val = m.group()
            if label == "EMAIL":
                if _SKIP_EXT.search(val):
                    continue
                try:
                    _, dom = val.rsplit("@", 1)
                    if dom.lower() not in _PERSONAL_DOMAINS:
                        continue
                except ValueError:
                    continue
            replacement = get_fake(val, label)
            modified = modified[:m.start()] + replacement + modified[m.end():]
    return modified

print("Processing DOCX...")
doc = Document(str(INPUT_PATH))

for para in doc.paragraphs:
    if para.text.strip():
        new_text = redact_string(para.text)
        if new_text != para.text:
            if para.runs:
                for r in para.runs[1:]:
                    r._element.getparent().remove(r._element)
                para.runs[0].text = new_text

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                new_text = redact_string(cell.text)
                if new_text != cell.text:
                    for p in cell.paragraphs:
                        if p.runs:
                            for r in p.runs[1:]:
                                r._element.getparent().remove(r._element)
                            p.runs[0].text = ""
                    if cell.paragraphs:
                        if cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].text = new_text
                        else:
                            cell.paragraphs[0].add_run(new_text)

doc.save(str(OUTPUT_PATH))
print(f"Output generated successfully: {OUTPUT_PATH}")
